import pyshark
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Hàm phân tích file PCAP
def analyze_pcap(file_path):
    capture = pyshark.FileCapture(file_path)
    protocol_stats = {}
    total_packets = 0
    total_size = 0
    tcp_packets = []
    udp_packets = []

    # Duyệt từng gói tin trong file PCAP
    for packet in capture:
        total_packets += 1
        packet_length = int(packet.length)
        total_size += packet_length

        protocol = packet.highest_layer
        if protocol in protocol_stats:
            protocol_stats[protocol]['count'] += 1
            protocol_stats[protocol]['size'] += packet_length
        else:
            protocol_stats[protocol] = {'count': 1, 'size': packet_length}

        if protocol == 'TCP':
            tcp_packets.append(packet)
        elif protocol == 'UDP':
            udp_packets.append(packet)

    return protocol_stats, total_packets, total_size, tcp_packets, udp_packets

# Hàm thêm đoạn văn bản định dạng vào tài liệu
def add_paragraph(doc, text, bold=False, align='left', font_size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if align == 'center' else WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style.font.size = Pt(font_size)

# Hàm phân tích và tạo nhận xét cho UDP
def generate_udp_analysis(udp_packets):
    total_udp = len(udp_packets)
    dns_count = 0
    dhcp_count = 0
    streaming_count = 0

    # Phân loại gói tin UDP
    for packet in udp_packets:
        if 'DNS' in packet.highest_layer:
            dns_count += 1
        elif 'DHCP' in packet.highest_layer:
            dhcp_count += 1
        else:
            streaming_count += 1  # Giả định phần còn lại là streaming
    
    # Tính toán phần trăm
    dns_percentage = (dns_count / total_udp) * 100 if total_udp > 0 else 0
    dhcp_percentage = (dhcp_count / total_udp) * 100 if total_udp > 0 else 0
    streaming_percentage = (streaming_count / total_udp) * 100 if total_udp > 0 else 0

    # Tạo nhận xét tự động dựa trên tỷ lệ phần trăm
    remarks = []
    remarks.append(f"Tổng số lượng gói tin UDP: {total_udp}")
    remarks.append(f"Tổng dung lượng: {sum(int(p.length) for p in udp_packets) / 1024:.2f} KB")
    remarks.append("Các dịch vụ UDP phổ biến:")

    if dns_count > 0:
        remarks.append(f"• DNS: {dns_count} gói tin ({dns_percentage:.2f}%), chủ yếu là các yêu cầu và phản hồi DNS.")
    if dhcp_count > 0:
        remarks.append(f"• DHCP: {dhcp_count} gói tin ({dhcp_percentage:.2f}%), liên quan đến việc cấp phát địa chỉ IP cho các thiết bị trong mạng.")
    if streaming_count > 0:
        remarks.append(f"• Streaming: {streaming_count} gói tin ({streaming_percentage:.2f}%), cho thấy có các dịch vụ streaming video hoặc audio.")

    remarks.append("Nhận xét:")
    if dns_count > 0 or dhcp_count > 0:
        remarks.append("• DNS và DHCP: Đây là các giao thức UDP tiêu chuẩn, hoạt động ổn định.")
    if streaming_count > 0:
        remarks.append("• Streaming: Tỷ lệ cao cho thấy có hoạt động streaming trong mạng, có thể dẫn đến tình trạng tắc nghẽn mạng nếu không quản lý tốt.")
    
    return remarks

# Hàm phân tích và tạo nhận xét cho TCP
def generate_tcp_remarks(tcp_stats, tcp_packet_count):
    remarks = []
    
    for flag, count in tcp_stats.items():
        percentage = (count / tcp_packet_count) * 100 if tcp_packet_count > 0 else 0
        if flag == 'SYN':
            remarks.append(f"• Gói SYN: Chiếm tỷ lệ {percentage:.2f}% trong tổng số gói tin TCP, cho thấy nhiều kết nối TCP mới được thiết lập.")
        elif flag == 'ACK':
            remarks.append(f"• Gói ACK: Chiếm {percentage:.2f}%, là dấu hiệu của các kết nối đang hoạt động bình thường với nhiều gói tin xác nhận (acknowledgement).")
        elif flag == 'FIN':
            remarks.append(f"• Gói FIN: Có tỷ lệ thấp ({percentage:.2f}%), cho thấy một số kết nối TCP đã được kết thúc.")
        elif flag == 'PSH':
            remarks.append(f"• Gói PSH: Chiếm {percentage:.2f}%, chỉ ra rằng có sự truyền tải dữ liệu giữa các máy.")

    return remarks

# Hàm tạo báo cáo
def generate_report(protocol_stats, total_packets, total_size, tcp_packets, udp_packets, output_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Set line spacing
    for style in doc.styles:
        if style.type == 1:  # Paragraph style
            style.paragraph_format.line_spacing = 1.5
    
    section = doc.sections[0]
    section.page_height = Inches(11.7)  # A4 size
    section.page_width = Inches(8.3)
    
    # Title và Overview
    doc.add_heading('Báo Cáo Phân Tích Dữ Liệu Mạng từ File PCAP', 0)
    doc.add_heading('1. Thông Tin Tổng Quan', level=1)
    add_paragraph(doc, f"Tên file: test_data.pcap", align='left', font_size=12)
    add_paragraph(doc, f"Thời gian bắt đầu phân tích: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align='left', font_size=12)
    add_paragraph(doc, f"Số lượng gói tin: {total_packets}", align='left', font_size=12)
    add_paragraph(doc, f"Tổng dung lượng: {total_size / 1024:.2f} KB", align='left', font_size=12)
    
    # Phân tích giao thức
    doc.add_heading('2. Phân Tích Chi Tiết Giao Thức', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Giao thức'
    hdr_cells[1].text = 'Số lượng gói tin'
    hdr_cells[2].text = 'Dung lượng (KB)'
    hdr_cells[3].text = 'Tỉ lệ (%)'

    for protocol, stats in protocol_stats.items():
        row_cells = table.add_row().cells
        row_cells[0].text = protocol
        row_cells[1].text = str(stats['count'])
        row_cells[2].text = f"{stats['size'] / 1024:.2f}"
        row_cells[3].text = f"{(stats['count'] / total_packets) * 100:.2f}"

    # Nhận xét giao thức
    doc.add_heading('Nhận xét:', level=2)
    for protocol, stats in protocol_stats.items():
        percentage = (stats['count'] / total_packets) * 100
        add_paragraph(doc, f"• {protocol}: Giao thức {protocol} chiếm tỷ lệ {percentage:.2f}% trong tổng số lượng gói tin.", font_size=12)
    
    # Phân tích TCP
    doc.add_heading('4. Phân Tích Lưu Lượng TCP', level=1)
    add_paragraph(doc, f'Tổng số lượng gói tin TCP: {len(tcp_packets)}', font_size=12)
    add_paragraph(doc, f'Tổng dung lượng: {sum(int(p.length) for p in tcp_packets) / 1024:.2f} KB', font_size=12)

    # Phân loại gói tin TCP
    doc.add_heading('Phân loại gói tin TCP:', level=2)
    table_tcp = doc.add_table(rows=1, cols=4)
    hdr_cells_tcp = table_tcp.rows[0].cells
    hdr_cells_tcp[0].text = 'Loại gói tin'
    hdr_cells_tcp[1].text = 'Số lượng'
    hdr_cells_tcp[2].text = 'Dung lượng (kB)'
    hdr_cells_tcp[3].text = 'Tỉ lệ (%)'

    # Phân tích các cờ TCP (SYN, ACK, FIN, PSH)
    tcp_stats = {
        'SYN': 0,
        'ACK': 0,
        'FIN': 0,
        'PSH': 0
    }
    for packet in tcp_packets:
        flags = packet.tcp.flags  # Lấy giá trị cờ TCP từ gói tin
        if 'S' in flags:  # SYN flag
            tcp_stats['SYN'] += 1
        if 'A' in flags:  # ACK flag
            tcp_stats['ACK'] += 1
        if 'F' in flags:  # FIN flag
            tcp_stats['FIN'] += 1
        if 'P' in flags:  # PSH flag
            tcp_stats['PSH'] += 1

    for flag, count in tcp_stats.items():
        row_cells_tcp = table_tcp.add_row().cells
        row_cells_tcp[0].text = flag
        row_cells_tcp[1].text = str(count)
        row_cells_tcp[2].text = f"{count * int(packet.length) / 1024:.2f}"
        row_cells_tcp[3].text = f"{(count / len(tcp_packets)) * 100:.2f}"

    # Tạo nhận xét TCP dựa trên phân tích
    doc.add_heading('Nhận xét:', level=2)
    tcp_remarks = generate_tcp_remarks(tcp_stats, len(tcp_packets))
    for remark in tcp_remarks:
        add_paragraph(doc, remark, font_size=12)

    # Phân tích UDP
    doc.add_heading('5. Phân Tích Lưu Lượng UDP', level=1)
    udp_analysis = generate_udp_analysis(udp_packets)
    for line in udp_analysis:
        add_paragraph(doc, line, font_size=12)

    # Lưu báo cáo
    doc.save(output_path)
    print(f"Report saved to {output_path}")

# Sử dụng ví dụ
file_path = 'test_data.pcapng'  # Thay bằng đường dẫn thực tế của bạn
output_report = 'report_udp.docx'

protocol_stats, total_packets, total_size, tcp_packets, udp_packets = analyze_pcap(file_path)
generate_report(protocol_stats, total_packets, total_size, tcp_packets, udp_packets, output_report)
