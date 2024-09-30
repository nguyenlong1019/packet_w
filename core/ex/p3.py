import pyshark
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Analyze the PCAP file
def analyze_pcap(file_path):
    capture = pyshark.FileCapture(file_path)
    protocol_stats = {}
    total_packets = 0
    total_size = 0
    tcp_packets = []
    udp_packets = []

    # Analyzing each packet
    for packet in capture:
        total_packets += 1
        packet_length = int(packet.length)
        total_size += packet_length

        protocol = packet.highest_layer
        if protocol in protocol_stats:
            protocol_stats[protocol]['count'] += 1
            protocol_stats[protocol]['size'] += packet_length
        else:
            protocol_stats[protocol] = {'count': 1, 'size': packet_length}

        if protocol == 'TCP':
            tcp_packets.append(packet)
        elif protocol == 'UDP':
            udp_packets.append(packet)

    return protocol_stats, total_packets, total_size, tcp_packets, udp_packets

# Add formatted text to the document
def add_paragraph(doc, text, bold=False, align='left', font_size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if align == 'center' else WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style.font.size = Pt(font_size)

# Generate remarks based on packet stats
def generate_tcp_remarks(tcp_stats, tcp_packet_count):
    remarks = []
    
    for flag, count in tcp_stats.items():
        percentage = (count / tcp_packet_count) * 100 if tcp_packet_count > 0 else 0
        if flag == 'SYN':
            remarks.append(f"• Gói SYN: Chiếm tỷ lệ {percentage:.2f}% trong tổng số gói tin TCP, cho thấy nhiều kết nối TCP mới được thiết lập.")
        elif flag == 'ACK':
            remarks.append(f"• Gói ACK: Chiếm {percentage:.2f}%, là dấu hiệu của các kết nối đang hoạt động bình thường với nhiều gói tin xác nhận (acknowledgement).")
        elif flag == 'FIN':
            remarks.append(f"• Gói FIN: Có tỷ lệ thấp ({percentage:.2f}%), cho thấy một số kết nối TCP đã được kết thúc.")
        elif flag == 'PSH':
            remarks.append(f"• Gói PSH: Chiếm {percentage:.2f}%, chỉ ra rằng có sự truyền tải dữ liệu giữa các máy.")

    return remarks

# Generate the report
def generate_report(protocol_stats, total_packets, total_size, tcp_packets, udp_packets, output_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Set line spacing
    for style in doc.styles:
        if style.type == 1:  # Paragraph style
            style.paragraph_format.line_spacing = 1.5
    
    section = doc.sections[0]
    section.page_height = Inches(11.7)  # A4 size
    section.page_width = Inches(8.3)
    
    # Title and overview
    doc.add_heading('Báo Cáo Phân Tích Dữ Liệu Mạng từ File PCAP', 0)
    doc.add_heading('1. Thông Tin Tổng Quan', level=1)
    add_paragraph(doc, f"Tên file: test_data.pcap", align='left', font_size=12)
    add_paragraph(doc, f"Thời gian bắt đầu phân tích: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align='left', font_size=12)
    add_paragraph(doc, f"Số lượng gói tin: {total_packets}", align='left', font_size=12)
    add_paragraph(doc, f"Tổng dung lượng: {total_size / 1024:.2f} KB", align='left', font_size=12)
    
    # Protocol analysis
    doc.add_heading('2. Phân Tích Chi Tiết Giao Thức', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Giao thức'
    hdr_cells[1].text = 'Số lượng gói tin'
    hdr_cells[2].text = 'Dung lượng (KB)'
    hdr_cells[3].text = 'Tỉ lệ (%)'

    for protocol, stats in protocol_stats.items():
        row_cells = table.add_row().cells
        row_cells[0].text = protocol
        row_cells[1].text = str(stats['count'])
        row_cells[2].text = f"{stats['size'] / 1024:.2f}"
        row_cells[3].text = f"{(stats['count'] / total_packets) * 100:.2f}"

    # Add observations for protocols dynamically
    doc.add_heading('Nhận xét:', level=2)
    for protocol, stats in protocol_stats.items():
        percentage = (stats['count'] / total_packets) * 100
        add_paragraph(doc, f"• {protocol}: Giao thức {protocol} chiếm tỷ lệ {percentage:.2f}% trong tổng số lượng gói tin.", font_size=12)
    
    # TCP Analysis
    doc.add_heading('4. Phân Tích Lưu Lượng TCP', level=1)
    add_paragraph(doc, f'Tổng số lượng gói tin TCP: {len(tcp_packets)}', font_size=12)
    add_paragraph(doc, f'Tổng dung lượng: {sum(int(p.length) for p in tcp_packets) / 1024:.2f} KB', font_size=12)

    # TCP Packet Breakdown (e.g., SYN, ACK, etc.)
    doc.add_heading('Phân loại gói tin TCP:', level=2)
    table_tcp = doc.add_table(rows=1, cols=4)
    hdr_cells_tcp = table_tcp.rows[0].cells
    hdr_cells_tcp[0].text = 'Loại gói tin'
    hdr_cells_tcp[1].text = 'Số lượng'
    hdr_cells_tcp[2].text = 'Dung lượng (kB)'
    hdr_cells_tcp[3].text = 'Tỉ lệ (%)'

    # Analyze TCP flags (SYN, ACK, FIN, etc.)
    tcp_stats = {
        'SYN': 0,
        'ACK': 0,
        'FIN': 0,
        'PSH': 0
    }
    for p in tcp_packets:
        print(p)
        break
    
    for packet in tcp_packets:
        
        if 'SYN' in packet.tcp.flags:
            tcp_stats['SYN'] += 1
        if 'ACK' in packet.tcp.flags:
            tcp_stats['ACK'] += 1
        if 'FIN' in packet.tcp.flags:
            tcp_stats['FIN'] += 1
        if 'PSH' in packet.tcp.flags:
            tcp_stats['PSH'] += 1

    for flag, count in tcp_stats.items():
        row_cells_tcp = table_tcp.add_row().cells
        row_cells_tcp[0].text = flag
        row_cells_tcp[1].text = str(count)
        row_cells_tcp[2].text = f"{count * int(packet.length) / 1024:.2f}"
        row_cells_tcp[3].text = f"{(count / len(tcp_packets)) * 100:.2f}"

    # Generate and add TCP remarks based on analysis
    doc.add_heading('Nhận xét:', level=2)
    # print(tcp_stats)
    tcp_remarks = generate_tcp_remarks(tcp_stats, len(tcp_packets))
    for remark in tcp_remarks:
        add_paragraph(doc, remark, font_size=12)

    # Similar analysis for UDP (you can expand this based on the same approach)
    
    doc.save(output_path)
    print(f"Report saved to {output_path}")

# Example Usage
file_path = 'test_data.pcapng'  # Replace with your actual file path
output_report = 'report2.docx'

protocol_stats, total_packets, total_size, tcp_packets, udp_packets = analyze_pcap(file_path)
generate_report(protocol_stats, total_packets, total_size, tcp_packets, udp_packets, output_report)
