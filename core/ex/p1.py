import pyshark
from datetime import datetime
from docx import Document

# Analyze the PCAP file
def analyze_pcap(file_path):
    capture = pyshark.FileCapture(file_path)
    protocol_stats = {}
    total_packets = 0
    total_size = 0

    # Analyzing each packet
    for packet in capture:
        total_packets += 1
        packet_length = int(packet.length)
        total_size += packet_length

        protocol = packet.highest_layer
        if protocol in protocol_stats:
            protocol_stats[protocol]['count'] += 1
            protocol_stats[protocol]['size'] += packet_length
        else:
            protocol_stats[protocol] = {'count': 1, 'size': packet_length}

    return protocol_stats, total_packets, total_size

# Generate the report
def generate_report(protocol_stats, total_packets, total_size, output_path):
    doc = Document()

    # Title and overview
    doc.add_heading('Báo Cáo Phân Tích Dữ Liệu Mạng từ File PCAP', 0)
    doc.add_heading('1. Thông Tin Tổng Quan', level=1)
    doc.add_paragraph(f"Tên file: test_data.pcap")
    doc.add_paragraph(f"Thời gian bắt đầu phân tích: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Số lượng gói tin: {total_packets}")
    doc.add_paragraph(f"Tổng dung lượng: {total_size / 1024:.2f} KB")
    
    # Protocol analysis
    doc.add_heading('2. Phân Tích Chi Tiết Giao Thức', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Giao thức'
    hdr_cells[1].text = 'Số lượng gói tin'
    hdr_cells[2].text = 'Dung lượng (KB)'
    hdr_cells[3].text = 'Tỉ lệ (%)'

    for protocol, stats in protocol_stats.items():
        row_cells = table.add_row().cells
        row_cells[0].text = protocol
        row_cells[1].text = str(stats['count'])
        row_cells[2].text = f"{stats['size'] / 1024:.2f}"
        row_cells[3].text = f"{(stats['count'] / total_packets) * 100:.2f}"

    doc.save(output_path)
    print(f"Report saved to {output_path}")

# Example Usage
file_path = 'test_data.pcapng'  # Replace with your actual file path
output_report = 'report.docx'

protocol_stats, total_packets, total_size = analyze_pcap(file_path)
generate_report(protocol_stats, total_packets, total_size, output_report)
